#!/usr/bin/env python3
"""生成标书智能体系统模拟参考材料"""
from pathlib import Path
from datetime import date, timedelta

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import Workbook

ROOT = Path(__file__).resolve().parent.parent / "sample_data"


def write_txt(path: Path, content: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content.strip() + "\n", encoding="utf-8")


def write_xlsx(path: Path, headers: list, rows: list):
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(path)


def make_template(path: Path):
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("铁路工程投标文件")
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph("项目名称：{{project_name}}")
    doc.add_paragraph("招标编号：{{tender_no}}")
    doc.add_paragraph("招标人：{{tenderer}}")
    doc.add_paragraph("投标人：{{bidder}}")
    doc.add_paragraph("投标总价：{{bid_amount}}")
    doc.add_paragraph("工期：{{duration}}")
    doc.add_paragraph("项目经理：{{project_manager}}")
    doc.add_paragraph("质保期：{{warranty_period}}")

    doc.add_heading("第一卷 商务标", level=1)
    doc.add_heading("第一章 投标函", level=2)
    doc.add_paragraph(
        "致 {{tenderer}}：我方已仔细研究了 {{project_name}}（招标编号：{{tender_no}}）"
        "招标文件的全部内容，愿意以人民币 {{bid_amount}} 的投标总价，"
        "按合同约定条件完成本项目，工期 {{duration}}，项目经理为 {{project_manager}}。"
    )
    doc.add_heading("第二章 法定代表人身份证明", level=2)
    doc.add_paragraph("（此处插入法定代表人身份证明材料）")
    doc.add_heading("第三章 授权委托书", level=2)
    doc.add_paragraph("（此处插入授权委托书）")
    doc.add_heading("第四章 投标保证金", level=2)
    doc.add_paragraph("（此处插入投标保证金凭证）")

    doc.add_heading("第二卷 技术标", level=1)
    doc.add_heading("第一章 工程概况", level=2)
    doc.add_paragraph("本工程为 {{project_name}}，招标人为 {{tenderer}}。")
    doc.add_heading("第二章 施工组织设计", level=2)
    doc.add_paragraph("【AI_GENERATED:施工组织设计】")
    doc.add_heading("第三章 人员配置", level=2)
    doc.add_paragraph("项目经理：{{project_manager}}")
    doc.add_paragraph("【AI_GENERATED:人员配置说明】")
    doc.add_heading("第四章 质量与安全保证措施", level=2)
    doc.add_paragraph("【AI_GENERATED:质量与安全保证措施】")

    doc.add_heading("第三卷 报价标", level=1)
    doc.add_heading("第一章 投标报价表", level=2)
    doc.add_paragraph("投标总价：{{bid_amount}}")
    doc.add_paragraph("【AI_GENERATED:分项报价说明】")

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "{{project_name}} 投标文件"
    footer = section.footer.paragraphs[0]
    footer.text = "第 页"

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def make_history_doc(path: Path, project: str, tenderer: str, amount: str, manager: str):
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("铁路工程投标文件（历史样例）")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph(f"项目名称：{project}")
    doc.add_paragraph("招标编号：TD-2024-DEMO-001")
    doc.add_paragraph(f"招标人：{tenderer}")
    doc.add_paragraph("投标人：XX铁路工程有限公司")
    doc.add_paragraph(f"投标总价：{amount}")
    doc.add_paragraph("工期：365日历天")
    doc.add_paragraph(f"项目经理：{manager}")
    doc.add_paragraph("质保期：24个月")

    doc.add_heading("第一卷 商务标", level=1)
    doc.add_heading("第一章 投标函", level=2)
    doc.add_paragraph(
        f"致 {tenderer}：我方愿以人民币 {amount} 的投标总价承接 {project}，"
        f"工期365日历天，项目经理为{manager}。"
    )
    doc.add_heading("第二章 法定代表人身份证明", level=2)
    doc.add_paragraph("已附法定代表人身份证明及授权委托书。")

    doc.add_heading("第二卷 技术标", level=1)
    doc.add_heading("第一章 工程概况", level=2)
    doc.add_paragraph(f"本工程为{project}，位于XX省XX市，建设内容含路基、桥梁、轨道等。")
    doc.add_heading("第二章 施工组织设计", level=2)
    doc.add_paragraph(
        "本工程采用分段流水施工组织方式，关键控制节点包括路基填筑、桥梁架设、"
        "轨道铺设及联调联试。施工期间严格执行铁路工程相关技术规范，"
        "配备专职安全员和质量员，实行日报周报制度。"
    )
    doc.add_heading("第三章 人员配置", level=2)
    doc.add_paragraph(f"项目经理：{manager}（一级建造师·铁路工程）")
    doc.add_paragraph("技术负责人：李某某（高级工程师）")
    doc.add_paragraph("安全负责人：王某某（安全员B证）")
    doc.add_heading("第四章 质量与安全保证措施", level=2)
    doc.add_paragraph(
        "建立三级质量检查体系，关键工序实行旁站监理配合；"
        "严格执行《铁路工程施工安全技术规程》，落实班前交底与隐患排查。"
    )

    doc.add_heading("第三卷 报价标", level=1)
    doc.add_heading("第一章 投标报价表", level=2)
    doc.add_paragraph(f"投标总价：{amount}")
    doc.add_paragraph("分项报价详见报价附件。")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def make_simple_docx(path: Path, title: str, body: str):
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    for line in body.strip().split("\n"):
        doc.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main():
    # 00
    write_txt(
        ROOT / "00_材料清单与说明" / "材料缺失说明.txt",
        "本次无缺失材料（模拟数据）。",
    )
    write_xlsx(
        ROOT / "00_材料清单与说明" / "材料提交清单.xlsx",
        ["序号", "目录路径", "文件名", "材料类型", "是否必填", "是否已提供", "备注"],
        [
            [1, "02_Word标书模板/标准投标Word模板/", "铁路工程投标模板_通用版.docx", "Word模板", "必填", "是", "含占位符"],
            [2, "03_历史投标文件/完整历史标书/", "历史标书_京沪高铁XX段_2024.docx", "历史标书", "必填", "是", ""],
            [3, "04_结构化字段与录入规范/", "标书关键字段清单.xlsx", "字段清单", "必填", "是", ""],
            [4, "05_条目完整性校验清单/", "商务标条目清单.xlsx", "条目清单", "必填", "是", ""],
            [5, "06_资质数据库参考材料/", "资质材料清单.xlsx", "资质清单", "必填", "是", ""],
            [6, "07_企业问答Chatbot参考/", "招标常见问题与标准答案.xlsx", "FAQ", "必填", "是", ""],
        ],
    )

    # 01
    write_txt(
        ROOT / "01_企业与项目基础信息" / "企业基本信息.txt",
        """
企业全称：XX铁路工程有限公司
企业简称：XX铁工
统一社会信用代码：91370000MA0000000X
注册地址：山东省济南市天桥区XX路XX号
办公地址：山东省济南市天桥区XX路XX号
法定代表人：张某某
注册资本：5000万元
成立日期：2010-05-18
联系电话：0531-88888888
企业邮箱：demo@xxrail.example.com
开户银行：中国邮政储蓄银行济南某支行
银行账号：937000000000000001
企业简介：本公司主要从事铁路、轨道交通工程施工，具备铁路工程施工总承包一级资质，近十年累计完成铁路工程业绩二十余项。
主营业务：铁路工程、轨道交通、站房建设
铁路/轨道交通相关资质概述：铁路工程施工总承包一级、安全生产许可证、ISO9001
""",
    )
    write_txt(
        ROOT / "01_企业与项目基础信息" / "典型投标项目说明.txt",
        """
1. 铁路工程：如京沪高铁XX段路基工程，招标方为铁路局集团公司，侧重施工组织与类似业绩。
2. 轨道交通：如城市地铁XX号线轨道工程，招标方为地铁公司，侧重盾构/轨道施工能力。
3. 站房改造：如XX站站房改扩建，招标方为地方铁路公司，侧重文明施工与工期保障。
""",
    )

    # 02
    make_template(ROOT / "02_Word标书模板" / "标准投标Word模板" / "铁路工程投标模板_通用版.docx")
    write_xlsx(
        ROOT / "02_Word标书模板" / "模板占位符说明.xlsx",
        ["占位符名称", "含义", "示例值", "出现位置"],
        [
            ["{{project_name}}", "项目名称", "XX铁路工程XX段", "封面、正文"],
            ["{{tender_no}}", "招标编号", "TD2026-001", "封面"],
            ["{{tenderer}}", "招标人", "XX铁路局集团有限公司", "封面、投标函"],
            ["{{bidder}}", "投标人", "XX铁路工程有限公司", "封面"],
            ["{{bid_amount}}", "投标总价", "人民币捌仟万元整", "报价部分"],
            ["{{duration}}", "工期", "365日历天", "技术部分"],
            ["{{project_manager}}", "项目经理", "张三", "人员部分"],
            ["{{warranty_period}}", "质保期", "24个月", "商务部分"],
        ],
    )
    write_txt(
        ROOT / "02_Word标书模板" / "模板格式规范说明.txt",
        """
正文：宋体 小四 1.5倍行距
一级标题：黑体 三号
二级标题：黑体 四号
页边距：上下2.54cm 左右3.17cm
""",
    )

    # 03
    histories = [
        ("历史标书_京沪高铁XX段_2024.docx", "京沪高铁XX段路基工程", "XX铁路局集团有限公司", "人民币捌仟万元整", "张三"),
        ("历史标书_地铁XX号线_2023.docx", "城市地铁XX号线轨道工程", "XX地铁集团有限公司", "人民币壹亿贰仟万元整", "李四"),
        ("历史标书_XX站房改造_2025.docx", "XX站站房改扩建工程", "地方铁路建设公司", "人民币叁仟伍佰万元整", "王五"),
    ]
    for name, project, tenderer, amount, manager in histories:
        make_history_doc(
            ROOT / "03_历史投标文件" / "完整历史标书" / name,
            project, tenderer, amount, manager,
        )
    write_txt(
        ROOT / "03_历史投标文件" / "历史标书章节结构说明.txt",
        """
第一卷 商务标
  第一章 投标函
  第二章 法定代表人身份证明
  第三章 授权委托书
  第四章 投标保证金
第二卷 技术标
  第一章 工程概况
  第二章 施工组织设计
  第三章 人员配置
  第四章 质量与安全保证措施
第三卷 报价标
  第一章 投标报价表
""",
    )
    write_txt(
        ROOT / "03_历史投标文件" / "AI生成内容风格参考.txt",
        """
语言风格：正式、严谨、符合铁路行业规范用语。
禁止使用：口语化表达、夸张修饰。
技术方案部分：需包含施工工艺、质量控制、安全措施等固定模块。
优秀段落样例：
本工程采用分段流水施工组织方式，关键控制节点包括路基填筑、桥梁架设、轨道铺设及联调联试。
""",
    )

    # 04
    write_xlsx(
        ROOT / "04_结构化字段与录入规范" / "标书关键字段清单.xlsx",
        ["字段名称", "字段英文名", "字段类型", "是否必填", "默认值", "校验规则", "所属模块", "备注"],
        [
            ["项目名称", "project_name", "文本", "是", "", "不超过200字", "基本信息", ""],
            ["招标编号", "tender_no", "文本", "是", "", "", "基本信息", ""],
            ["招标人", "tenderer", "文本", "是", "", "", "基本信息", ""],
            ["投标人", "bidder", "文本", "是", "XX铁路工程有限公司", "", "基本信息", ""],
            ["项目类型", "project_type", "下拉选项", "是", "", "", "基本信息", ""],
            ["投标总价", "bid_amount", "金额", "是", "", "大于0", "报价部分", ""],
            ["工期", "duration", "文本", "是", "365日历天", "", "技术部分", ""],
            ["项目经理", "project_manager", "人员", "是", "", "", "人员部分", ""],
            ["质保期", "warranty_period", "文本", "否", "24个月", "", "商务部分", ""],
        ],
    )
    write_xlsx(
        ROOT / "04_结构化字段与录入规范" / "下拉选项枚举值.xlsx",
        ["字段名称", "可选值（多个值用中文分号分隔）"],
        [
            ["项目类型", "铁路工程;轨道交通;站房建设;线路维护;其他"],
            ["招标方式", "公开招标;邀请招标;竞争性谈判"],
            ["资质等级", "特级;一级;二级"],
        ],
    )
    write_txt(
        ROOT / "04_结构化字段与录入规范" / "铁路行业专用术语表.txt",
        """
TB 10401 《铁路工程施工质量验收标准》
TB 10301 《铁路工程设计技术规范》
联调联试：工程完工后对信号、供电、轨道等系统进行综合调试。
""",
    )

    # 05
    write_xlsx(
        ROOT / "05_条目完整性校验清单" / "商务标条目清单.xlsx",
        ["序号", "条目名称", "是否必含", "对应章节", "备注"],
        [
            [1, "投标函", "必含", "第一章", ""],
            [2, "法定代表人身份证明", "必含", "第二章", ""],
            [3, "授权委托书", "必含", "第三章", ""],
            [4, "投标保证金凭证", "必含", "第四章", ""],
            [5, "联合体协议书", "条件必含", "第五章", "仅联合体投标时"],
        ],
    )
    write_xlsx(
        ROOT / "05_条目完整性校验清单" / "技术标条目清单.xlsx",
        ["序号", "条目名称", "是否必含", "对应章节", "备注"],
        [
            [1, "工程概况", "必含", "第一章", ""],
            [2, "施工组织设计", "必含", "第二章", ""],
            [3, "人员配置", "必含", "第三章", ""],
            [4, "质量与安全保证措施", "必含", "第四章", ""],
            [5, "工期保障措施", "必含", "第五章", ""],
        ],
    )
    write_xlsx(
        ROOT / "05_条目完整性校验清单" / "报价标条目清单.xlsx",
        ["序号", "条目名称", "是否必含", "对应章节", "备注"],
        [
            [1, "投标报价表", "必含", "第一章", ""],
            [2, "分项报价", "必含", "第二章", ""],
            [3, "主要材料价格表", "建议", "第三章", ""],
        ],
    )
    write_xlsx(
        ROOT / "05_条目完整性校验清单" / "资质材料必附清单.xlsx",
        ["序号", "资质名称", "所属分类", "适用项目类型", "是否必附", "备注"],
        [
            [1, "营业执照", "企业资质包", "全部", "必附", ""],
            [2, "铁路工程专业承包资质", "企业资质包", "铁路工程", "必附", ""],
            [3, "近3年类似业绩", "业绩包", "全部", "必附", "至少3个"],
            [4, "项目经理一级建造师证", "人员信息包", "全部", "必附", ""],
            [5, "近三年审计报告", "财务信息包", "全部", "必附", ""],
        ],
    )

    # 06 资质
    qual_root = ROOT / "06_资质数据库参考材料"
    today = date.today()
    quals = [
        ("01_企业资质包", "营业执照", "企业资质_营业执照_长期.docx", True, "市场监督管理局"),
        ("01_企业资质包", "铁路工程施工总承包一级", "企业资质_铁路总承包一级_20301231.docx", False, "住建部"),
        ("01_企业资质包", "安全生产许可证", "企业资质_安全生产许可证_20281231.docx", False, "应急管理部门"),
        ("02_业绩包", "京沪高铁XX段合同", "业绩_京沪高铁XX段合同_2024.docx", True, "XX铁路局"),
        ("02_业绩包", "地铁XX号线中标通知书", "业绩_地铁XX号线中标通知_2023.docx", True, "XX地铁集团"),
        ("03_人员信息包", "项目经理简历及建造师证", "人员_项目经理张三_建造师.docx", False, "住建部"),
        ("03_人员信息包", "技术负责人职称证", "人员_技术负责人李某某_职称.docx", False, "人社部门"),
        ("04_财务信息包", "近三年审计报告", "财务_近三年审计报告_2025.docx", True, "会计师事务所"),
        ("05_信誉与法律包", "无行贿犯罪证明", "信誉_无行贿证明_2026.docx", False, "检察机关"),
        ("06_技术方案包", "施工组织设计模板", "技术_施工组织设计模板.docx", True, "企业内部"),
        ("07_商务文件包", "投标函模板", "商务_投标函模板.docx", True, "企业内部"),
        ("07_商务文件包", "授权委托书模板", "商务_授权委托书模板.docx", True, "企业内部"),
    ]
    rows = []
    for i, (cat, name, fname, long_term, issuer) in enumerate(quals, 1):
        expire = "" if long_term else (today + timedelta(days=365 * (2 if "2030" in fname else 1))).isoformat()
        start = (today - timedelta(days=365)).isoformat()
        make_simple_docx(
            qual_root / cat / fname,
            name,
            f"材料名称：{name}\n颁发机构：{issuer}\n有效期起：{start}\n有效期止：{expire or '长期有效'}\n（模拟材料，仅供系统开发测试）",
        )
        rows.append([
            i, cat.split("_", 1)[1], name, fname, start, expire or "长期",
            issuer, "docx", name, "是" if long_term else "否", "模拟材料",
        ])
    write_xlsx(
        qual_root / "资质材料清单.xlsx",
        ["序号", "分类", "材料名称", "文件名", "有效期起", "有效期止", "颁发机构", "文件类型", "关键词标签", "是否长期有效", "备注"],
        rows,
    )
    write_txt(
        qual_root / "资质插入规则说明.txt",
        """
企业资质章节 → 营业执照 + 铁路工程施工总承包资质 + 安全生产许可证
人员章节 → 项目经理建造师证 + 简历
业绩章节 → 类似项目合同 + 中标通知书
商务部分 → 投标函模板 + 授权委托书模板
""",
    )

    # 07 FAQ
    write_xlsx(
        ROOT / "07_企业问答Chatbot参考" / "招标常见问题与标准答案.xlsx",
        ["序号", "问题类别", "问题内容", "标准答案", "答案来源材料", "备注"],
        [
            [1, "企业资质类", "贵公司是否具备铁路工程施工总承包一级资质？",
             "是，我公司具备铁路工程施工总承包一级资质。", "企业资质_铁路总承包一级_20301231.docx", ""],
            [2, "业绩类", "近5年是否有类似铁路项目业绩？",
             "有，我公司近5年完成京沪高铁XX段路基工程等多项类似业绩。", "业绩_京沪高铁XX段合同_2024.docx", ""],
            [3, "人员类", "拟任项目经理是否具备一级建造师（铁路）？",
             "是，拟任项目经理张三持有一级建造师（铁路工程）注册证书。", "人员_项目经理张三_建造师.docx", ""],
            [4, "财务类", "是否可提供近三年审计报告？",
             "可以，我公司可提供近三年完整审计报告。", "财务_近三年审计报告_2025.docx", ""],
            [5, "信誉类", "近三年是否有行贿犯罪记录？",
             "无，我公司可提供无行贿犯罪记录证明。", "信誉_无行贿证明_2026.docx", ""],
            [6, "技术类", "施工组织设计包含哪些主要内容？",
             "包含工程概况、施工部署、进度计划、质量与安全保证措施等。", "技术_施工组织设计模板.docx", ""],
            [7, "企业资质类", "是否持有安全生产许可证？",
             "是，我公司持有有效期内的安全生产许可证。", "企业资质_安全生产许可证_20281231.docx", ""],
            [8, "业绩类", "是否有城市轨道交通项目经验？",
             "有，我公司完成城市地铁XX号线轨道工程。", "业绩_地铁XX号线中标通知_2023.docx", ""],
        ],
    )
    write_txt(
        ROOT / "07_企业问答Chatbot参考" / "招标文件常见要求摘录.txt",
        """
常见评分点：类似业绩、项目经理资质、施工组织设计完整性、财务状况、信誉情况。
常见硬性要求：铁路工程施工总承包一级、安全生产许可证、无重大安全事故声明。
""",
    )

    print(f"模拟材料已生成至: {ROOT}")


if __name__ == "__main__":
    main()
